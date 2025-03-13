import os
from dotenv import load_dotenv
from datetime import datetime, timedelta
import streamlit as st
import pandas as pd
from io import BytesIO

# Importar los agentes del sistema
from agents.consultor import TaskManager
from agents.data_wrangler import DataWrangler
from agents.meta_specialist import MetaSpecialist
from agents.account_manager import AccountManager

# Importar la librería python-docx para generar documentos de Word
from docx import Document

def recommendations_to_markdown(recommendations: dict) -> str:
    """
    Convierte el diccionario de recomendaciones en formato Markdown.
    
    Se espera que las recomendaciones sean un diccionario que contenga,
    opcionalmente, una clave "recommendations". Para cada categoría, se generan
    encabezados y se listan las recomendaciones. Si el valor asociado a una clave no es un diccionario,
    se agrega como texto simple.
    
    Parámetros:
    - recommendations (dict): Diccionario con las recomendaciones.
    
    Retorna:
    - str: Recomendaciones formateadas en Markdown.
    """
    md = ""
    rec = recommendations.get("recommendations", recommendations)
    for category in rec:
        md += f"## {str(category).capitalize()}\n\n"
        subdict = rec[category]
        # Si subdict es un diccionario, iterar sobre sus items
        if isinstance(subdict, dict):
            for key, value in subdict.items():
                md += f"### {str(key).replace('_', ' ').capitalize()}\n"
                if isinstance(value, dict):
                    for subkey, subvalue in value.items():
                        md += f"- **{str(subkey).replace('_', ' ').capitalize()}:** {subvalue}\n"
                else:
                    md += f"- {value}\n"
                md += "\n"
        else:
            # Si no es un diccionario, lo agregamos como una entrada de lista
            md += f"- {subdict}\n\n"
    return md


def generate_docx(informe: str, recomendaciones: dict) -> BytesIO:
    """
    Genera un documento de Word (DOCX) con el informe y las recomendaciones en formato estructurado.
    
    - Se utiliza la librería python-docx para crear el documento.
    - Se agregan encabezados y listas para estructurar la información.
    
    Parámetros:
    - informe (str): Texto del informe generado por el Meta Specialist.
    - recomendaciones (dict): Diccionario con recomendaciones estructuradas.
    
    Retorna:
    - BytesIO: Documento de Word en memoria listo para descarga.
    """
    doc = Document()
    # Agregar el título del informe
    doc.add_heading("Informe de Meta Specialist", level=1)
    doc.add_paragraph(informe)
    
    # Agregar el título de recomendaciones
    doc.add_heading("Recomendaciones del Account Manager", level=1)
    md_reco = recommendations_to_markdown(recomendaciones)
    for line in md_reco.splitlines():
        # Formatear como títulos y listas según corresponda
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph(line)
    
    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def run_pipeline(user_prompt: str):
    """
    Ejecuta el flujo de procesamiento de datos y generación de informes.
    
    Pasos:
    1. Carga las credenciales y configura BigQuery.
    2. Genera el input estructurado con TaskManager.
    3. Extrae los datos de BigQuery con DataWrangler.
    4. Analiza los datos con MetaSpecialist.
    5. Genera recomendaciones con AccountManager.
    
    Parámetros:
    - user_prompt (str): Solicitud del usuario en lenguaje natural.
    
    Retorna:
    - tuple (informe, recomendaciones): Texto del informe y recomendaciones generadas.
    """
    # Cargar variables de entorno y credenciales de BigQuery
    load_dotenv()
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "D:/jordiquiroga.com/Scripts/multiagentes_reporting/credentials.json"

    st.header("Generación del Input estructurado")
    
    # 1. Generar input estructurado con el TaskManager
    tm = TaskManager(verbose=True)
    input_json = tm.generar_inputs(user_prompt)
    st.write("**Inputs generados por el TaskManager:**")
    st.json(input_json)

    # Encapsular el input si no tiene la clave 'solicitud' ni 'request'
    if "solicitud" not in input_json and "request" not in input_json:
        input_json = {"solicitud": input_json}

    # 2. Ajustar las métricas en caso de que no sean las esperadas
    if "request" in input_json:
        req = input_json["request"]
    elif "solicitud" in input_json:
        req = input_json["solicitud"]
    else:
        req = input_json

    # Si las métricas incluyen solo "conversions" y "lead", o "conversions_lead",
    # se reemplazan por un conjunto de métricas más amplio por defecto.
    if ("metrics" not in req) or (set(req.get("metrics", [])) == {"conversions", "lead"} or "conversions_lead" in req.get("metrics", [])):
        req["metrics"] = ["impresiones", "clics", "gasto", "conversions"]

    st.write("**Input modificado para DataWrangler:**")
    st.json(input_json)

    # 3. Extraer datos desde BigQuery usando el DataWrangler
    dw = DataWrangler()
    df = dw.extraer_datos(input_json)
    
    if df is None or df.empty:
        st.warning("No se han extraído datos con la configuración actual. Intentando con un período calculado para los últimos 90 días...")
        
        # Calcular dinámicamente el período de los últimos 90 días
        today = datetime.today().date()
        start_date = today - timedelta(days=90)
        new_time_period = {"start_date": str(start_date), "end_date": str(today)}

        if "solicitud" in input_json:
            input_json["solicitud"]["time_period"] = new_time_period
        elif "request" in input_json:
            input_json["request"]["time_period"] = new_time_period
        else:
            input_json["time_period"] = new_time_period

        st.write("**Nuevo input con el período actualizado:**")
        st.json(input_json)

        # Reintentar la extracción de datos con el nuevo período
        df = dw.extraer_datos(input_json)

    if df is None or df.empty:
        st.error("No se han extraído datos. Verifica los filtros y el período.")
        return None, None
    else:
        st.write("**DataFrame final extraído:**")
        st.dataframe(df)

    # 4. Analizar los datos con el Meta Specialist
    ms = MetaSpecialist()
    informe = ms.analizar(df)
    st.header("Informe de Meta Specialist")
    st.text(informe)

    # 5. Generar recomendaciones con el Account Manager
    am = AccountManager()
    recomendaciones = am.generar_recomendaciones(informe)
    st.header("Recomendaciones del Account Manager")
    st.json(recomendaciones)
    
    return informe, recomendaciones

def main():
    """
    Función principal para ejecutar la aplicación en Streamlit.
    
    - Presenta un campo de entrada para la solicitud del usuario.
    - Ejecuta el pipeline de generación de informes al presionar el botón.
    - Permite descargar el informe en formato Word.
    """
    st.title("Sistema de Agentes para Reporting de Campañas de Facebook Ads")
    st.write("Ingrese la solicitud del usuario:")

    # Campo de texto para la solicitud del usuario
    user_prompt = st.text_area(
        "Solicitud del usuario", 
        "Quiero un informe del rendimiento de mis campañas de Facebook Ads en los últimos 90 días, "
        "filtrado por device_platform 'mobile_app' y que incluya conversiones de lead."
    )

    if st.button("Ejecutar Pipeline"):
        informe, recomendaciones = run_pipeline(user_prompt)
        if informe and recomendaciones:
            st.success("Pipeline ejecutado con éxito!")
            # Botón para descargar el informe en formato Word
            doc_buffer = generate_docx(informe, recomendaciones)
            st.download_button(
                label="Exportar Informe a Word",
                data=doc_buffer,
                file_name="informe.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()







